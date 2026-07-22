"""Document ingestion and structure-aware chunking utilities."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterator


@dataclass(frozen=True)
class DocumentPart:
    """Text from one logical part of a source document."""

    text: str
    source_file: str
    page_number: int | None = None
    section_title: str | None = None


_HEADING = re.compile(r"^(?:#{1,6}\s+|\d+(?:\.\d+)*[.)]?\s+|[A-Z][A-Z\s]{3,}:?$)")


def ingest_file(path: str | Path, *, ocr: bool = True) -> list[DocumentPart]:
    """Extract text and page/section provenance from a supported local file.

    OCR is used only when a PDF page has no extractable text. Image OCR requires
    the system ``tesseract`` executable in addition to the Python dependency.
    """
    source = Path(path)
    if not source.is_file():
        raise FileNotFoundError(source)
    suffix = source.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _with_sections(source.read_text(encoding="utf-8"), source.name)
    if suffix == ".docx":
        return _ingest_docx(source)
    if suffix == ".pdf":
        return _ingest_pdf(source, ocr=ocr)
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return [DocumentPart(_ocr_image(source), source.name, page_number=1)]
    raise ValueError(f"Unsupported document type: {source.suffix}")


def _ingest_pdf(path: Path, *, ocr: bool) -> list[DocumentPart]:
    try:
        from pypdf import PdfReader
    except ImportError as error:  # pragma: no cover - dependency guard
        raise RuntimeError("PDF support requires pypdf; install requirements.txt") from error
    parts: list[DocumentPart] = []
    for page_number, page in enumerate(PdfReader(str(path)).pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text and ocr:
            text = _ocr_pdf_page(path, page_number)
        if text:
            parts.extend(DocumentPart(part.text, path.name, page_number, part.section_title)
                         for part in _with_sections(text, path.name))
    return parts


def _ingest_docx(path: Path) -> list[DocumentPart]:
    try:
        from docx import Document
    except ImportError as error:  # pragma: no cover - dependency guard
        raise RuntimeError("DOCX support requires python-docx; install requirements.txt") from error
    document = Document(path)
    parts: list[DocumentPart] = []
    section: str | None = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            section = text
        else:
            parts.append(DocumentPart(text, path.name, section_title=section))
    # Tables preserve rows and field/value pairs as a single retrieval unit.
    for table in document.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if any(rows):
            parts.append(DocumentPart("\n".join(rows), path.name, section_title=section or "Table"))
    return parts


def _ocr_image(path: Path) -> str:
    try:
        from PIL import Image
        import pytesseract
    except ImportError as error:  # pragma: no cover - dependency guard
        raise RuntimeError("Image OCR requires Pillow and pytesseract; install requirements.txt") from error
    return pytesseract.image_to_string(Image.open(path)).strip()


def _ocr_pdf_page(path: Path, page_number: int) -> str:
    # TODO: Add a PDF renderer dependency (for example PyMuPDF) to OCR scanned
    # PDF pages directly. Current fallback provides a clear actionable error.
    raise RuntimeError(
        f"Page {page_number} of {path.name} is image-only. Convert it to an image "
        "and use image OCR, or install a future PDF-rendering OCR extra."
    )


def _with_sections(text: str, source_file: str) -> list[DocumentPart]:
    section: str | None = None
    parts: list[DocumentPart] = []
    buffer: list[str] = []
    def flush() -> None:
        if buffer:
            parts.append(DocumentPart("\n".join(buffer).strip(), source_file, section_title=section))
            buffer.clear()
    for block in re.split(r"\n\s*\n", text.strip()):
        block = block.strip()
        if not block:
            continue
        first = block.splitlines()[0].strip()
        if _HEADING.match(first) and len(first) < 140:
            flush()
            section = re.sub(r"^#+\s*", "", first).rstrip(":")
            remaining = "\n".join(block.splitlines()[1:]).strip()
            if remaining:
                buffer.append(remaining)
        else:
            buffer.append(block)
    flush()
    return parts or [DocumentPart(text.strip(), source_file)]


def chunk_parts(parts: list[DocumentPart], max_words: int = 180) -> Iterator[DocumentPart]:
    """Chunk by paragraph/row boundaries, splitting only oversized units."""
    for part in parts:
        units = [item.strip() for item in part.text.split("\n") if item.strip()]
        current: list[str] = []
        count = 0
        for unit in units or [part.text]:
            words = unit.split()
            if current and count + len(words) > max_words:
                yield DocumentPart("\n".join(current), part.source_file, part.page_number, part.section_title)
                current, count = [], 0
            while len(words) > max_words:
                if current:
                    yield DocumentPart("\n".join(current), part.source_file, part.page_number, part.section_title)
                    current, count = [], 0
                yield DocumentPart(" ".join(words[:max_words]), part.source_file, part.page_number, part.section_title)
                words = words[max_words:]
            current.append(" ".join(words))
            count += len(words)
        if current:
            yield DocumentPart("\n".join(current), part.source_file, part.page_number, part.section_title)
